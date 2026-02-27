"""
Generate Word documents from markdown content for AK Dental partnership docs.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_table_from_md(doc, header_row, data_rows):
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        hdr_cells[i].text = cell_text.strip()
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        # Shade header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text.strip()

    doc.add_paragraph()  # spacing after table


def add_bold_paragraph(doc, line, style="Normal"):
    """Add a paragraph that supports **bold** inline markup."""
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


def parse_and_add_line(doc, line, in_table_buffer, table_buffer):
    """Returns (consumed_table, new_in_table_buffer, new_table_buffer)."""
    # Heading 1
    if re.match(r"^# (.+)$", line):
        if in_table_buffer:
            flush_table(doc, table_buffer)
            table_buffer.clear()
            in_table_buffer = False
        text = re.match(r"^# (.+)$", line).group(1)
        para = doc.add_heading(level=1)
        para.clear()
        add_inline_bold(para, text)
        return False, in_table_buffer, table_buffer

    # Heading 2
    if re.match(r"^## (.+)$", line):
        if in_table_buffer:
            flush_table(doc, table_buffer)
            table_buffer.clear()
            in_table_buffer = False
        text = re.match(r"^## (.+)$", line).group(1)
        para = doc.add_heading(level=2)
        para.clear()
        add_inline_bold(para, text)
        return False, in_table_buffer, table_buffer

    # Heading 3
    if re.match(r"^### (.+)$", line):
        if in_table_buffer:
            flush_table(doc, table_buffer)
            table_buffer.clear()
            in_table_buffer = False
        text = re.match(r"^### (.+)$", line).group(1)
        para = doc.add_heading(level=3)
        para.clear()
        add_inline_bold(para, text)
        return False, in_table_buffer, table_buffer

    # Horizontal rule
    if re.match(r"^---+$", line.strip()):
        if in_table_buffer:
            flush_table(doc, table_buffer)
            table_buffer.clear()
            in_table_buffer = False
        return False, False, table_buffer

    # Table row
    if line.strip().startswith("|"):
        cells = [c for c in line.strip().split("|") if c != ""]
        # Skip separator rows like |---|---|
        if all(re.match(r"^[\s\-:]+$", c) for c in cells):
            return False, in_table_buffer, table_buffer
        table_buffer.append(cells)
        return False, True, table_buffer

    # If we were in a table and hit a non-table line, flush it
    if in_table_buffer:
        flush_table(doc, table_buffer)
        table_buffer.clear()
        in_table_buffer = False

    # Checkbox item  - [ ]
    if re.match(r"^\s*- \[ \] (.+)$", line):
        text = re.match(r"^\s*- \[ \] (.+)$", line).group(1)
        para = doc.add_paragraph(style="List Bullet")
        para.clear()
        para.add_run("☐  ")
        add_inline_bold_to_para(para, text)
        return False, False, table_buffer

    # Italic/note line  *...*
    if re.match(r"^\*(.+)\*$", line.strip()):
        text = re.match(r"^\*(.+)\*$", line.strip()).group(1)
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        return False, False, table_buffer

    # Numbered list
    if re.match(r"^\d+\. (.+)$", line.strip()):
        text = re.match(r"^\d+\. (.+)$", line.strip()).group(1)
        para = doc.add_paragraph(style="List Number")
        para.clear()
        add_inline_bold_to_para(para, text)
        return False, False, table_buffer

    # Bullet list (- item)
    if re.match(r"^- (.+)$", line.strip()):
        text = re.match(r"^- (.+)$", line.strip()).group(1)
        para = doc.add_paragraph(style="List Bullet")
        para.clear()
        add_inline_bold_to_para(para, text)
        return False, False, table_buffer

    # Empty line
    if line.strip() == "":
        return False, False, table_buffer

    # Normal paragraph
    para = doc.add_paragraph()
    add_inline_bold_to_para(para, line.strip())
    return False, False, table_buffer


def add_inline_bold(para, text):
    """Add runs with **bold** support to an existing paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                para.add_run(part)


def add_inline_bold_to_para(para, text):
    add_inline_bold(para, text)


def flush_table(doc, table_buffer):
    if not table_buffer:
        return
    header = table_buffer[0]
    data = table_buffer[1:] if len(table_buffer) > 1 else []
    add_table_from_md(doc, header, data)


def process_markdown(doc, markdown_text):
    lines = markdown_text.splitlines()
    in_table_buffer = False
    table_buffer = []

    for line in lines:
        _, in_table_buffer, table_buffer = parse_and_add_line(
            doc, line, in_table_buffer, table_buffer
        )

    # Flush any remaining table
    if in_table_buffer and table_buffer:
        flush_table(doc, table_buffer)


# ---------------------------------------------------------------------------
# Document 1: PARTNERSHIP-OUTLINE.docx
# ---------------------------------------------------------------------------
PARTNERSHIP_OUTLINE_MD = r"""# Alex Chireau DDS — Dental Engine Partnership Outline
**Date:** February 2026
**Prepared by:** Brad Palubicki, NuStack Digital Ventures
**Status:** In Discussion

---

## The Opportunity

NuStack Digital Ventures is building **Dental Engine** — a SaaS platform purpose-built for dental practices. It automates scheduling, clinical notes, billing, patient communications, and integrates with existing practice management systems (Dentrix, Eaglesoft, Open Dental).

Alex Chireau DDS and AK Ultimate Dental are positioned to be **Partner Clinic #1** — the founding test clinic that shapes the product and shares in its future.

---

## What Alex Is Being Offered

### Investment
- **$20,000** investment into Dental Engine (not NuStack Digital Ventures as a whole)
- Brad investing $20,000 + Ken investing $20,000 = $60,000 total launch capital
- Proceeds fund: product completion, vendor integrations, sales infrastructure

### Ownership
- Alex receives **equity ownership in Dental Engine only**
- Not in NuStack Digital Ventures, not in other NuStack engines
- Ownership percentage to be determined in operating agreement (proportional to investment)

### Profitability
- Alex shares in **Dental Engine profits only**
- As Dental Engine scales to other practices, Alex's stake generates passive income
- Revenue model: monthly SaaS subscriptions from dental practices nationwide

### What Alex Does NOT Have To Do
Alex does NOT need to:
- Build anything
- Manage staff or operations
- Handle sales or marketing
- Work more than a few hours per quarter

---

## Alex's Role

### Spokesperson (Minimal Time)
- Appear in 2-3 short video testimonials (filmed at his office)
- Allow use of name/credentials in marketing materials: "Built with and trusted by AK Ultimate Dental, Las Vegas"
- Available for occasional reference calls with prospective dental clients
- Estimated time: **2-4 hours total per quarter**

### Test Clinic (Active Partnership)
- AK Ultimate Dental runs the live Dental Engine platform
- Alex/staff provides feedback on what works and what doesn't
- Early access to all new features before general release
- Brad/NuStack handles all implementation, support, and maintenance
- Staff training provided at no cost

---

## What NuStack Delivers to Alex

1. **AK Ultimate Dental website** — fully built, hosted, live at custom domain
2. **Dental Engine platform** — full access for AK Ultimate Dental, no monthly fee
3. **Integrations**: Dentrix, credit card processing, electronic billing, payroll
4. **Patient communications**: automated appointment reminders, recall campaigns
5. **Analytics dashboard**: production reports, patient retention metrics
6. **Ongoing development**: new features added continuously

---

## Capital Allocation ($60K Total)

| Use | Amount |
|-----|--------|
| Product completion + QA | $20,000 |
| Vendor integrations (Dentrix API, etc.) | $10,000 |
| Marketing + first 10 client acquisition | $15,000 |
| Legal (operating agreements, IP assignment) | $5,000 |
| Operating reserve | $10,000 |

---

## Timeline to Next Milestone

| Milestone | Target |
|-----------|--------|
| Partnership agreement signed | March 2026 |
| Vendor info gathered from Alex | March 2026 |
| Dentrix API integration complete | April 2026 |
| AK Dental live on full platform | April 2026 |
| First external dental client onboarded | May 2026 |
| Revenue sharing begins | Upon first client |

---

## Information Needed From Alex

See VENDOR-INFO-REQUEST for the full list.

**Key items:**
- Dentrix account/subscription details (for API integration)
- Current credit card processor (name, contact)
- Payroll provider
- Electronic billing / clearinghouse currently used
- NPI number(s) for the practice
- Tax ID / EIN for the practice entity

---

## Next Steps to Move Forward

1. Alex confirms interest in $20K investment tier
2. Brad sends operating agreement draft (attorney review)
3. Alex provides vendor information sheet
4. NuStack begins API integrations
5. AK Dental migrates to new domain + Dental Engine platform
6. Partnership formally launched

---

*This document is a working outline only. All terms subject to formal operating agreement.*
"""

# ---------------------------------------------------------------------------
# Document 2: VENDOR-INFO-REQUEST.docx
# ---------------------------------------------------------------------------
VENDOR_INFO_MD = r"""# Vendor Information Request — AK Ultimate Dental
**For:** Alex Chireau DDS
**Purpose:** API integration planning for Dental Engine platform
**Date:** February 2026

---

## Why We Need This

To fully integrate Dental Engine with your existing practice systems, we need account details and contact information for each vendor. We will handle all API setup, integration, and ongoing maintenance — we just need the baseline info to get started.

Nothing will be changed or disrupted at your practice without your approval.

---

## Practice Management System (PMS)

### Dentrix (Henry Schein)
- [ ] Dentrix version currently installed (e.g., Dentrix G7, G8)
- [ ] Henry Schein account number
- [ ] Contact name at your Dentrix account rep
- [ ] Is Dentrix Ascend (cloud) or server-based?
- [ ] Do you currently use the Dentrix patient portal?
- [ ] Are you on a support/maintenance plan? (needed for API access)

*Note: Dentrix API integration requires an active support plan and Henry Schein partner approval. We will initiate the partner application process.*

---

## Credit Card Processing

- [ ] Current processor name (e.g., Stripe, Square, Heartland, Clover, TransFirst, Worldpay)
- [ ] Contact name / account rep
- [ ] Account number or merchant ID
- [ ] Do you currently use a patient-facing payment portal?
- [ ] Monthly processing volume (approximate)

*Note: We will integrate your existing processor into the patient billing flow, or recommend a switch to Square if that's preferred for better rates/API support.*

---

## Electronic Billing / Insurance Clearinghouse

- [ ] Clearinghouse currently used (e.g., Availity, Waystar, Change Healthcare, Dental Xchange, ClaimRemedi)
- [ ] Account login email
- [ ] Contact name at clearinghouse
- [ ] Do you process claims directly or through a billing service?
- [ ] Name of billing service/coordinator if outsourced

*Note: We will build automated claim submission and ERA posting into the platform.*

---

## Payroll

- [ ] Payroll provider (e.g., Gusto, ADP, Paychex, QuickBooks Payroll)
- [ ] Account login email or admin contact
- [ ] Number of employees (approximate)
- [ ] Pay frequency (weekly, biweekly, semi-monthly)

*Note: Payroll integration is Phase 2. We need the vendor now to plan the integration.*

---

## Other Systems (If Applicable)

### Recall / Patient Communication
- [ ] Any current recall system? (e.g., Solutionreach, Lighthouse 360, Weave, RevenueWell, Demandforce)
- [ ] If yes: account name, contact

### Online Scheduling
- [ ] Do you currently use any online booking tool? (e.g., NexHealth, LocalMed, ZocDoc)

### Reviews / Reputation
- [ ] Google Business Profile — are you the owner/manager? Can you add Brad as manager?
- [ ] Yelp Business account access?

### Phone / Communication
- [ ] Current phone system (landline, VoIP, etc.)
- [ ] Provider name (e.g., RingCentral, Weave, 8x8, Vonage, standard POTS)

---

## Practice Identifiers (Required for Billing Integration)

- [ ] Practice NPI (Type 2 — Organization NPI)
- [ ] Provider NPI(s) — Alex's individual NPI
- [ ] Tax ID / EIN for the practice entity
- [ ] Practice legal name (as it appears on insurance contracts)
- [ ] State dental license number

---

## How to Submit

You can simply reply to Brad's email with the items you know, or fill this in as best you can. We will follow up on anything missing. Nothing is urgent — we can gather this over 1-2 conversations.

---

*All information provided will be kept strictly confidential and used only for the purpose of platform integration.*
"""

# ---------------------------------------------------------------------------
# Document 3: WEBSITE-STATUS.docx
# ---------------------------------------------------------------------------
WEBSITE_STATUS_MD = r"""# AK Ultimate Dental — Website & Platform Status
**As of:** February 2026
**Prepared by:** Brad Palubicki, NuStack Digital Ventures

---

## Website: What's Live

The new AK Ultimate Dental website has been fully built and is live at the Vercel preview URL. It is ready to go live on the production domain (akultimatedental.com) whenever Alex gives the green light.

### Pages Built

| Page | URL | Status |
|------|-----|--------|
| Homepage | / | Live |
| About | /about | Live |
| Services (all) | /services | Live |
| Individual service pages | /services/[service] | Live (20+ services) |
| Technology | /technology | Live |
| Blog | /blog | Live (12 articles) |
| Reviews | /reviews | Live |
| Contact | /contact | Live |
| Appointment Request | /appointment | Live |
| FAQs | /faqs | Live |
| Patient Portal | /patient-portal | Live |
| Dentist Las Vegas | /dentist-las-vegas | Live |
| Careers | /careers | Live |
| Dr. Scott Miller Retirement | /dr-scott-miller-retirement | Live |

### Technical Stack
- **Framework**: Next.js 16 (App Router) — fast, SEO-optimized
- **Hosting**: Vercel (CDN, automatic HTTPS, global edge network)
- **Images**: Professional dental photography via Next.js Image optimization
- **SEO**: Full schema markup (LocalBusiness, MedicalProcedure, FAQ, Article, BreadcrumbList)
- **Performance**: Lighthouse score 90+ target
- **Mobile**: Fully responsive for all devices

---

## SEO: What's Built In

### Structured Data (Schema.org)
Every page has proper JSON-LD markup for:
- Local Business (practice name, address, phone, hours, geo coordinates)
- Aggregate Rating (Google review count + score)
- Medical Procedure (for each service page)
- FAQ (for FAQ pages)
- Article (for blog posts)
- Breadcrumb navigation

### Metadata
- Unique title tags and meta descriptions for every page
- Open Graph tags for social sharing and AI search (ChatGPT, Perplexity)
- Canonical URLs
- Sitemap (auto-generated)

### Content
- 12 blog articles written targeting high-value Las Vegas dental search terms
- Service pages with full procedure descriptions
- FAQs targeting "People Also Ask" Google results
- 301 redirects from all old site URLs (so existing rankings transfer)

---

## Dental Engine: What's Built

The backend SaaS platform powering the practice operations:

| Feature | Status |
|---------|--------|
| Patient intake forms | Complete |
| Online appointment requests | Complete |
| Appointment types + scheduling | Complete |
| Provider availability management | Complete |
| Clinical notes (AI-assisted) | Complete |
| Dental charting | Complete |
| CDT code library (31 codes) | Complete |
| Lab case tracking | Complete |
| Patient consent forms | Complete |
| Recall system (automated) | Complete |
| Assessment tools | Complete |
| Dashboard (admin + provider views) | Complete |
| Role-based access control | Complete |
| Outreach / patient communications | Complete |

### Still To Build (Phase 2)
- Dentrix PMS integration (needs API credentials)
- Insurance billing / clearinghouse integration
- Credit card processing integration
- Payroll integration
- Patient portal full features

---

## Domain Cutover Plan

When ready to go live on akultimatedental.com:

1. Add Brad as contributor in Vercel (or we handle DNS change together)
2. Update Vercel project domain settings
3. Change DNS at domain registrar to point to Vercel
4. SSL certificate auto-provisions (takes < 5 minutes)
5. Old site 301 redirects fire automatically — all existing Google rankings transfer
6. Google Search Console: submit new sitemap

**Estimated downtime:** Near zero (< 60 seconds for DNS propagation edge cases)

---

## What's NOT Changing

- Phone number stays the same
- Address stays the same
- All existing patient reviews stay visible on Google/Yelp
- The brand name and visual identity are preserved and enhanced
- Existing URL paths from the old site redirect automatically to new pages

---

*The site is ready. The platform is ready. We're waiting on your go-ahead.*
"""

# ---------------------------------------------------------------------------
# Document 4: EMAIL-DRAFT.docx  (with real values substituted)
# ---------------------------------------------------------------------------
EMAIL_DRAFT_MD = r"""# Email Draft — Alex Chireau DDS
**To:** akultimatedentalpc@yahoo.com
**From:** brad@nustack.digital
**Subject:** Ready to Move Forward — Partnership Overview + Next Steps

---

Hi Alex,

Hope things are going well at the practice. I wanted to reach out because I think we're at the point where it makes sense to move forward, and I want to make sure I give you everything you need to make that decision.

Since we last spoke, a lot has been accomplished. The AK Ultimate Dental website is fully built — every page, every service, the blog, reviews integration, the works. The backend platform (Dental Engine) is also built out with scheduling, clinical notes, charting, patient communications, and more. I've attached a full status document so you can see exactly where everything stands.

---

**The Opportunity**

Here's what I'm proposing, and I want to be straightforward about the structure so there's no ambiguity:

- The investment is **$20,000** into Dental Engine specifically — not NuStack as a whole
- Ken and I are each putting in $20,000 as well — so we're at $60K total to get this to market
- In return, you'd receive **equity ownership in Dental Engine** and share in its **profitability as it grows**
- This is the ground floor — we're at the test clinic stage now, and the upside is significant as we scale to other dental practices

---

**What You'd Actually Have to Do**

I know this was your question — and the honest answer is: not much.

1. **Be the face of it.** A couple short video testimonials, permission to use your name and credentials in our marketing ("Built with AK Ultimate Dental, Las Vegas"), and occasional availability for a reference call with a prospective client. That's maybe 2-4 hours a quarter total.

2. **Let us use AK Dental as the live platform.** Your practice runs on Dental Engine. Your staff uses the tools. When something needs improvement, you tell us. We fix it. You and your team never have to manage any of the technology — that's entirely on us.

That's it. No sales. No management. No operational headaches. You run your practice, and the platform runs in the background.

---

**What I Need From You to Get Started**

To complete the platform integrations for AK Dental, I need some vendor information — specifically around your practice management system (I believe you're using Dentrix), credit card processing, and billing setup. I've attached a simple checklist. You don't need to fill it all out at once — even a partial response gets us moving.

In parallel, I'll have the operating agreement drafted for attorney review. I want everything properly documented so this is a clean, protected relationship for both of us.

---

**Attached Documents**

1. **Partnership Outline** — Full structure of the deal, Alex's role, capital allocation, timeline
2. **Vendor Information Request** — What we need from you to complete the integrations
3. **Website & Platform Status** — Full rundown of everything that's been built

---

**Next Steps**

1. Let me know if the $20K investment and the structure makes sense for you
2. Send back the vendor info sheet (or just reply with what you know)
3. I'll get the operating agreement drafted and over to you
4. We'll schedule a call to walk through anything you want to review in detail

Alex, I genuinely believe this is a great opportunity — and more importantly, I think you're the right partner for this. Your credentials, your reputation in Las Vegas, and the practice you've built give Dental Engine instant credibility in the market. That's worth a lot.

Let me know your thoughts. Happy to jump on a call anytime.

Brad
brad@nustack.digital
(608) 215-3940

---

*Attachments: Partnership-Outline.pdf | Vendor-Info-Request.pdf | Website-Platform-Status.pdf*
"""


def create_doc(markdown_text, filename):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    process_markdown(doc, markdown_text)

    out_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    create_doc(PARTNERSHIP_OUTLINE_MD, "PARTNERSHIP-OUTLINE.docx")
    create_doc(VENDOR_INFO_MD, "VENDOR-INFO-REQUEST.docx")
    create_doc(WEBSITE_STATUS_MD, "WEBSITE-STATUS.docx")
    create_doc(EMAIL_DRAFT_MD, "EMAIL-DRAFT.docx")
    print("All 4 documents created successfully.")
